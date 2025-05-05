import fitz  # PyMuPDF for PDF processing
from docx import Document  # For DOCX processing
import os
from PIL import Image
import io
import re


class ImageExtractor:
    def extract(self, file_path):
        print(f"Extracting images from: {file_path}")
        ext = os.path.splitext(file_path)[1].lower()

        if ext == '.pdf':
            return self._extract_from_pdf(file_path)
        elif ext == '.docx':
            return self._extract_from_docx(file_path)
        else:
            print("Unsupported file format for image extraction.")
            return []

    def _extract_from_pdf(self, file_path):
        images = []
        doc = fitz.open(file_path)

        # First, get all page texts to help with context lookup
        page_texts = [page.get_text() for page in doc]

        for page_num, page in enumerate(doc, start=1):
            page_text = page_texts[page_num-1]
            
            # Extract figure/table captions before processing images
            captions = self._extract_captions(page_text)
            
            # Process images on the page
            image_list = page.get_images(full=True)
            for img_index, img in enumerate(image_list):
                try:
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    img_ext = base_image["ext"]
                    
                    # Skip SVG images as they often cause issues
                    if img_ext.lower() == 'svg':
                        continue
                        
                    image = Image.open(io.BytesIO(image_bytes))

                    # Skip very small images (likely icons or bullets)
                    if image.width < 100 or image.height < 100:
                        continue
                        
                    # Skip excessively wide images that might be horizontal rules
                    if image.width > image.height * 10:
                        continue

                    os.makedirs("extracted/images", exist_ok=True)
                    image_path = f"extracted/images/pdf_page{page_num}_img{img_index + 1}.{img_ext}"
                    image.save(image_path)

                    # Find the best caption for this image
                    best_caption = self._find_best_caption(captions, img_index)
                    
                    # If no specific caption, extract surrounding text for context
                    if not best_caption:
                        # Get text in a window around the image's position
                        rect = page.get_image_bbox(xref)
                        if rect:
                            # Get text above and below the image
                            text_above = page.get_text("text", clip=(rect[0], rect[1]-200, rect[2], rect[1]))
                            text_below = page.get_text("text", clip=(rect[0], rect[3], rect[2], rect[3]+200))
                            surrounding_text = (text_above + " " + text_below).strip()
                            
                            # If surrounding text is found, use it as caption
                            if surrounding_text:
                                best_caption = surrounding_text[:200]  # Limit context length
                    
                    # If still no caption, use page text
                    if not best_caption:
                        # Extract keywords from page text for context
                        best_caption = self._extract_keywords(page_text)
                    
                    images.append({
                        "path": image_path,
                        "page": page_num,
                        "context": best_caption,
                        "size": image.size
                    })

                except Exception as e:
                    print(f"Failed to extract image on page {page_num}: {e}")
                    continue

        return images

    def _extract_from_docx(self, file_path):
        images = []
        doc = Document(file_path)
        
        # Extract all paragraphs for context
        paragraphs = [p.text for p in doc.paragraphs]
        
        # Identify caption paragraphs
        captions = []
        for i, para in enumerate(paragraphs):
            if re.search(r'(Figure|Fig\.?|Table)\s*\d+', para, re.IGNORECASE):
                captions.append((i, para))

        for i, rel in enumerate(doc.part._rels.values()):
            if "image" in rel.target_ref:
                try:
                    image_data = rel.target_part.blob
                    image = Image.open(io.BytesIO(image_data))
                    
                    # Skip tiny images (likely icons or bullets)
                    if image.width < 100 or image.height < 100:
                        continue

                    os.makedirs("extracted/images", exist_ok=True)
                    image_path = f"extracted/images/docx_image{i + 1}.png"
                    image.save(image_path)

                    # Find the best context for this image
                    context = ""
                    
                    # First try to find the paragraph that references this image
                    for para_idx, para in enumerate(doc.paragraphs):
                        if rel.target_ref in para._element.xml:
                            # Found the paragraph with this image
                            # Get context from surrounding paragraphs
                            start_idx = max(0, para_idx - 1)
                            end_idx = min(len(doc.paragraphs), para_idx + 2)
                            
                            context_paragraphs = [doc.paragraphs[j].text for j in range(start_idx, end_idx)]
                            context = " ".join(context_paragraphs)
                            break
                    
                    # If no paragraph references this image directly, look for nearby captions
                    if not context:
                        # Find caption closest to this image's reference
                        for para_idx, para in enumerate(doc.paragraphs):
                            if rel.target_ref in para._element.xml:
                                # Look for the closest caption
                                closest_caption = None
                                min_distance = float('inf')
                                
                                for caption_idx, caption_text in captions:
                                    distance = abs(caption_idx - para_idx)
                                    if distance < min_distance:
                                        min_distance = distance
                                        closest_caption = caption_text
                                
                                if closest_caption and min_distance <= 3:  # Within 3 paragraphs
                                    context = closest_caption
                                break
                    
                    # If still no context, extract keywords from nearby text
                    if not context:
                        # Get some surrounding paragraphs for context
                        relevant_paragraphs = []
                        for para_idx, para in enumerate(doc.paragraphs):
                            if rel.target_ref in para._element.xml:
                                # Get 3 paragraphs before and after
                                start_idx = max(0, para_idx - 3)
                                end_idx = min(len(doc.paragraphs), para_idx + 4)
                                relevant_paragraphs = [doc.paragraphs[j].text for j in range(start_idx, end_idx)]
                                break
                        
                        context = self._extract_keywords(" ".join(relevant_paragraphs))
                    
                    images.append({
                        "path": image_path,
                        "index": i + 1,
                        "context": context,
                        "size": image.size
                    })

                except Exception as e:
                    print(f"Failed to extract image from DOCX: {e}")
                    continue

        return images

    def _extract_captions(self, text):
        """Extract figure and table captions from text"""
        captions = []
        
        # Match common caption patterns
        caption_patterns = [
            r'(Figure|Fig\.?)\s*(\d+)[\.:]?\s*([^\n\.]+)',
            r'(Table)\s*(\d+)[\.:]?\s*([^\n\.]+)'
        ]
        
        for pattern in caption_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                caption_type = match.group(1)
                caption_num = match.group(2)
                caption_text = match.group(3).strip()
                full_caption = f"{caption_type} {caption_num}: {caption_text}"
                captions.append((caption_num, full_caption))
        
        return captions

    def _find_best_caption(self, captions, img_index):
        """Find the most likely caption for an image"""
        if not captions:
            return ""
            
        # First try to match by index
        for caption_num, caption_text in captions:
            if str(img_index + 1) == caption_num:
                return caption_text
        
        # If no match by index, return the closest caption
        # Assuming images and captions are ordered similarly
        if img_index < len(captions):
            return captions[img_index][1]
        elif captions:
            return captions[-1][1]
            
        return ""

    def _extract_keywords(self, text):
        """Extract important keywords from text for context"""
        # Simple keyword extraction - get first 150 chars
        if len(text) <= 150:
            return text
            
        # Look for sentences with important keywords
        important_keywords = ["figure", "table", "chart", "graph", "image", "diagram", "illustration", 
                              "photo", "plot", "map", "screenshot", "picture"]
        
        sentences = re.split(r'[.!?]\s+', text)
        for sentence in sentences:
            if any(keyword in sentence.lower() for keyword in important_keywords):
                return sentence[:200]  # Limit length
        
        # If no important sentences found, return first 150 chars
        return text[:150]